import os
import shelve
import subprocess
import sys
from datetime import date

from pyperclip import paste
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt
from docx.shared import Inches


def main():
    # os.chdir(r"C:\Users\Franklin\CS50-Final-Project")

    if (
        3 < len(sys.argv) < 6
        and sys.argv[1].lower() == "letter"
        and sys.argv[2].lower() == "to"
    ):
        create_document()

    elif len(sys.argv) == 2 and sys.argv[1].lower() == "save":
        print("\n", save_address(), "\n", sep="")

    elif len(sys.argv) == 2 and sys.argv[1].lower() == "list":
        list_all_address()

    else:
        sys.exit(
            """\
    project.py - Create MS Word templates for formal letters 
    
    Usage:  python project.py letter to <company name> - Create a template with the address of the specified company.
            
            python project.py save - Store the address of a specific company in the database
            
            python project.py list - List all company names in the database in alphabetical order
            """
        )


def create_document():
    """
    This function calls all the functions allocated to the various components of a formal letter
    and creates a word template with customised styles and font for each type of letter

    :return: name_of_document.docx
    """
    doc = Document()

    # Set custom style, font, font size and alignment for entire document
    base_style = doc.styles["Normal"]
    base_style.font.name = "Times New Roman"
    base_style.font.size = Pt(12)
    base_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Set custom margins
    section = doc.sections[0]
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # Add Draft
    p_draft = doc.add_paragraph(draft())
    p_draft.runs[0].bold = True
    p_draft.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add Reference
    p_reference = doc.add_paragraph(reference())
    p_reference.runs[0].bold = True
    p_reference.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add Date
    p_date_doc = doc.add_paragraph(date_doc())
    p_date_doc.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Add Address
    p_address = doc.add_paragraph(company_address())
    p_address.runs[0].bold = True
    p_address.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add Greeting or salutation
    para_greet = doc.add_paragraph(greeting())
    para_greet.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add Title
    p_title = doc.add_paragraph(title())
    p_title.runs[0].underline = True
    p_title.runs[0].bold = True
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add body of letter
    p_body = doc.add_paragraph(body())
    p_body.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # Add valediction
    p_valediction = doc.add_paragraph(valediction())
    p_valediction.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add name of writer
    p_writer_name = doc.add_paragraph(writer_name())
    p_writer_name.runs[0].bold = True
    p_writer_name.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Save Document
    number = 1
    while True:
        doc_name = (
            f"{sys.argv[1]}_{sys.argv[2]}_{'_'.join(sys.argv[3:])}_{str(number)}.docx"
        )
        if not os.path.exists(doc_name):
            break
        number += 1

    doc.save(doc_name)
    print("\n\t" + doc_name + " has been created...\n")
    
    file_path = os.path.join("/Users/franklinaryee/CS50-Final-Project", doc_name)

    
    # Path to Microsoft Word application
    word_app_path = "/Applications/Microsoft Word.app"
    

    # Open document
    subprocess.run(["open", "-a", word_app_path, file_path], check=True)


def draft():
    """
    :return: A string containing `DRAFT`
    :rtype: str
    """
    return "DRAFT"


def reference():
    """
    This function returns the standard reference format when writing letters to companies

    :return: A sting of the reference format of the letter
    :rtype: str
    """
    ref_year = str(date.today().year)[2:]
    return "EC/LCLP/EMO/" + ref_year + "/0.."


def date_doc():
    """
    This function returns the current date of letter

    :return: A sting of the date of letter
    :rtype: str
    """
    return date.today().strftime("%B %d, %Y")


def company_address():
    """
    This function returns the address of the company passed in the command-line `(sys.argv[3])`
    or calls the `save_address` if the company name cannot be found in the `database`

    :rtype: str
    """
    with shelve.open("data_base") as sfile:
        company_name = " ".join(sys.argv[3:]).lower()

        if company_name in sfile:
            return sfile[company_name]

    print(
        f"Address of {company_name} cannot be found. "
        f"Thus, follow the next steps to add the address to the database..."
    )
    return save_address()


def greeting():
    """
    :return: A sting containing `Dear Sir/Madam,`
    :rtype: str
    """
    return "Dear Sir/Madam,"


def title():
    """
    This function returns `ENTER TITLE HERE` which indicates where the yet to be determined title will be typed.

    :return: A string of the title of the letter
    :rtype: str
    """
    return "ENTER TITLE HERE:"


def body():
    """
    This function returns `...` indicating where the body of the letter should be typed

    :return: A sting of the body of the letter
    :rtype: str
    """
    return "..."


def valediction():
    """
    This function returns the complimentary close or valediction
    and separates it from the body of the letter with whitespaces

    :return: A string of valediction of the letter
    :rtype: str
    """
    return ("\n" * 10) + "Yours faithfully,\n"


def writer_name():
    """
    :return: A string of the writer's name and position
    :rtype: str
    """
    return "Ing. Oscar Amonoo-Neizer\n(Executive Secretary)"


def save_address():
    """
    This function asks the user the `name` and `address` of the company when called.
    See `company_address` for further details.

    If the input of `name_of_company` is empty, the program will raise an error
    (You must type the name of company)

    If the input of `address_of_company` is empty, the address will be taken from
    the first item on the clipboard otherwise, the address with have to be typed
    manually in the format
    `(Position of recipient, Company name, House/Building/Flat number, Street, City)`

    :rtype: str
    """
    with shelve.open("data_base", writeback=True) as sfile:
        if len(sys.argv) < 3:
            name_of_company = (
                input("Enter the name of the company whose address you want to save: ")
                .lower()
                .strip()
            )
            if not name_of_company:
                raise ValueError("Name of company cannot be empty. Please try again.")
        else:
            name_of_company = " ".join(sys.argv[3:]).lower()
        address_of_company = (
            input(
                "\nAddress of company (e.g. The Managing Director, Volta River Authority, Electro-Volta House, "
                "Accra)\nEnter address: "
            )
            .title()
            .strip()
        )
        if not address_of_company:
            address_of_company = paste().strip().title()
            sfile[name_of_company] = address_of_company.replace("\r", "")
        else:
            sfile[name_of_company] = address_of_company.replace(", ", "\n")
        address_of_company = sfile[name_of_company]

    print(f"\nAddress of {name_of_company} has been saved.")
    return address_of_company


def list_all_address():
    """
    This function returns a list of all company names in the `database`

    :return: # . Company name
    :rtype: str
    """
    with shelve.open("data_base") as sfile:
        list_of_dicts = [
            {key: value} for key, value in zip(list(sfile.keys()), list(sfile.values()))
        ]
        if not list_of_dicts:
            print("The database is empty!\n")
        else:
            for _, key in enumerate(sorted(sfile.keys()), 1):
                print(f"{_}. {key}\n")


if __name__ == "__main__":
    main()
